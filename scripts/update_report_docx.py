from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Báo Cáo Đồ Án Truy Hồi Thông Tin - Nhóm 5.docx"
OUT_DIR = ROOT / "output" / "doc"
BACKUP_PATH = OUT_DIR / "Bao_Cao_Do_An_Truy_Hoi_Thong_Tin_Nhom_5.before-update.docx"
UPDATED_COPY_PATH = OUT_DIR / "Bao_Cao_Do_An_Truy_Hoi_Thong_Tin_Nhom_5.updated.docx"
FALLBACK_COPY_PATH = OUT_DIR / "Bao_Cao_Do_An_Truy_Hoi_Thong_Tin_Nhom_5.updated.v2.docx"

SectionContent = list[tuple[str, str]]

SECTION_CONTENT: dict[str, SectionContent] = {}
REFERENCE_LINES: list[str] = []
HEADING_RENAMES: dict[str, tuple[str, str]] = {
    "GIỚI THIỆU": ("Heading 1", "CHƯƠNG 1. GIỚI THIỆU"),
    "Giới thiệu đề tài": ("Heading 2", "1.1. Giới thiệu đề tài"),
    "Mục tiêu và nội dung thực hiện": ("Heading 2", "1.2. Mục tiêu và nội dung thực hiện"),
    "Giới hạn đề tài": ("Heading 2", "1.3. Giới hạn đề tài"),
    "Bố cục báo cáo": ("Heading 2", "1.4. Bố cục báo cáo"),
    "Cơ sở lý thuyết": ("Heading 1", "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT"),
    "Các khái niệm cơ bản": ("Heading 2", "2.1. Các khái niệm cơ bản"),
    "Từ Boolean retrieval đến ranked retrieval": ("Heading 3", "2.1.1. Từ Boolean retrieval đến ranked retrieval"),
    "Mô hình không gian vector": ("Heading 3", "2.1.2. Mô hình không gian vector"),
    "Trọng số term: TF, IDF và TF-IDF": ("Heading 3", "2.1.3. Trọng số term: TF, IDF và TF-IDF"),
    "Cosine similarity và độ liên quan hình học": ("Heading 3", "2.1.4. Cosine similarity và độ liên quan hình học"),
    "Mô hình xác suất, PRP và Binary Independence Model": ("Heading 3", "2.1.5. Mô hình xác suất, PRP và Binary Independence Model"),
    "BM25, TF saturation và ý nghĩa xác suất của IDF": ("Heading 3", "2.1.6. BM25, TF saturation và ý nghĩa xác suất của IDF"),
    "Dense retrieval, hybrid retrieval, RAG và LLM-only": ("Heading 3", "2.1.7. Dense retrieval, hybrid retrieval, RAG và LLM-only"),
    "Các chỉ số đánh giá và relevance feedback": ("Heading 3", "2.1.8. Các chỉ số đánh giá và relevance feedback"),
    "Các công trình nghiên cứu liên quan (ít nhất 5 công trình mới nhất)": ("Heading 2", "2.2. Các công trình nghiên cứu liên quan (ít nhất 5 công trình mới nhất)"),
    "Thiết kế triển khai hệ thống": ("Heading 1", "CHƯƠNG 3. THIẾT KẾ TRIỂN KHAI HỆ THỐNG"),
    "Tổng quan phương pháp hiện thực": ("Heading 2", "3.1. Tổng quan phương pháp hiện thực"),
    "Mô tả hệ thống (System Overview)": ("Heading 2", "3.2. Mô tả hệ thống (System Overview)"),
    "Pipeline tổng thể": ("Heading 3", "3.2.1. Pipeline tổng thể"),
    "Mô tả từng module": ("Heading 3", "3.2.2. Mô tả từng module"),
    "Công cụ hiện thực": ("Heading 2", "3.3. Công cụ hiện thực"),
    "Tập dữ liệu": ("Heading 2", "3.4. Tập dữ liệu"),
    "Quy trình xây dựng Knowledge Base trong đề tài": ("Heading 3", "3.4.1. Quy trình xây dựng Knowledge Base trong đề tài"),
    "Mô hình truy hồi và chiến lược trả lời": ("Heading 2", "3.5. Mô hình truy hồi và chiến lược trả lời"),
    "THỰC NGHIỆM": ("Heading 1", "CHƯƠNG 4. THỰC NGHIỆM"),
    "Thiết lập các cài đặt mô hình": ("Heading 2", "4.1. Thiết lập các cài đặt mô hình"),
    "Đánh giá 2 phương pháp": ("Heading 2", "4.2. Đánh giá 2 phương pháp"),
    "Kết quả thực nghiệm": ("Heading 2", "4.3. Kết quả thực nghiệm"),
    "4.3.1. Kết quả retrieval": ("Heading 3", "4.3.1. Kết quả retrieval"),
    "Kết quả answer-level evaluation": ("Heading 3", "4.3.2. Kết quả answer-level evaluation"),
    "4.3.3. Pairwise comparison": ("Heading 3", "4.3.3. Pairwise comparison"),
    "4.3.4. Truy vấn minh họa trên demo app": ("Heading 3", "4.3.4. Truy vấn minh họa trên demo app"),
    "Phân tích và thảo luận kết quả": ("Heading 2", "4.4. Phân tích và thảo luận kết quả"),
    "KẾT LUẬN": ("Heading 1", "CHƯƠNG 5. KẾT LUẬN"),
    "Nhận xét kết quả đề tài": ("Heading 2", "5.1. Nhận xét kết quả đề tài"),
    "Ưu – nhược điểm của đề tài": ("Heading 2", "5.2. Ưu – nhược điểm của đề tài"),
    "Hướng phát triển": ("Heading 2", "5.3. Hướng phát triển"),
    "TÀI LIỆU THAM KHẢO": ("Heading 1", "TÀI LIỆU THAM KHẢO"),
}

SECTION_CONTENT.update(
    {
        "Giới thiệu đề tài": [
            (
                "Normal",
                "Trong các hệ thống hỏi đáp hiện đại, hai hướng tiếp cận phổ biến là trả lời trực tiếp bằng tri thức tham số của mô hình ngôn ngữ lớn và trả lời dựa trên ngữ cảnh được truy hồi từ kho tài liệu ngoài. Hướng thứ nhất triển khai nhanh và có khả năng diễn đạt linh hoạt, nhưng khó kiểm soát nguồn gốc thông tin. Hướng thứ hai chậm và phức tạp hơn, song có ưu thế về khả năng bám bằng chứng và giải thích vì sao câu trả lời được đưa ra.",
            ),
            (
                "Normal",
                "Đối với nhóm câu hỏi sức khỏe và dinh dưỡng, yêu cầu về tính đúng, tính thận trọng và khả năng truy vết nguồn là rất quan trọng. Một câu trả lời nghe tự nhiên nhưng không có căn cứ rõ ràng có thể làm người dùng hiểu sai vấn đề. Vì vậy, đề tài lựa chọn bài toán nutrition-focused health question answering như một môi trường phù hợp để so sánh giữa cách trả lời grounded và cách trả lời closed-book.",
            ),
            (
                "List Paragraph",
                "Hybrid RAG: truy hồi tài liệu bằng tín hiệu lexical và semantic, sau đó sinh câu trả lời có trích dẫn.",
            ),
            (
                "List Paragraph",
                "Contextual Hybrid RAG: mở rộng grounded branch bằng contextual chunking, contextual summary và contextual retrieval.",
            ),
            (
                "List Paragraph",
                "LLM-only: trả lời trực tiếp bằng mô hình ngôn ngữ mà không sử dụng ngữ cảnh truy hồi ngoài.",
            ),
            (
                "Normal",
                "Repo hiện tại được tổ chức như một hệ thống local-first gồm bốn pha chính là chuẩn bị dữ liệu, indexing, retrieval và evaluation, kèm theo một demo app để quan sát trực tiếp pipeline. Dữ liệu đầu vào được giới hạn trong miền dinh dưỡng và thông tin sức khỏe phổ thông, sử dụng NFCorpus làm backbone benchmark, đồng thời bổ sung MedlinePlus, FDA và PubMed để tăng độ phong phú của evidence.",
            ),
            (
                "Normal",
                "Đề tài không hướng tới xây dựng một sản phẩm tư vấn y khoa hoàn chỉnh. Mục tiêu chính là tạo ra một môi trường thực nghiệm có thể tái chạy, có số liệu định lượng và có giao diện minh họa, từ đó đánh giá rõ hơn khi nào retrieval giúp ích cho câu trả lời và khi nào mô hình trả lời trực tiếp vẫn còn hạn chế.",
            ),
            ("Caption", "Hình 1.1. Mô hình tổng quan của dự án"),
        ],
        "Mục tiêu và nội dung thực hiện": [
            (
                "Normal",
                "Mục tiêu tổng quát của đề tài là xây dựng một hệ thống thực nghiệm để so sánh hai hướng trả lời câu hỏi sức khỏe định hướng dinh dưỡng: grounded answering dựa trên truy hồi bằng chứng và non-grounded answering dựa trên tri thức tham số của mô hình. Việc so sánh không chỉ dừng ở chất lượng câu trả lời cuối, mà còn bao gồm toàn bộ chuỗi xử lý từ dữ liệu, chỉ mục, truy hồi đến đánh giá.",
            ),
            (
                "Normal",
                "Mục tiêu thứ nhất là xây dựng knowledge base phục vụ truy hồi. Nhóm thu thập và chuẩn hóa dữ liệu từ NFCorpus, MedlinePlus, FDA và PubMed, sau đó đưa về schema thống nhất để có thể index trong PostgreSQL. Ở bước này, hệ thống không chỉ lưu văn bản gốc mà còn bổ sung các trường metadata, thông tin chunk và contextual summary để hỗ trợ nhiều kiểu truy hồi khác nhau.",
            ),
            (
                "Normal",
                "Mục tiêu thứ hai là hiện thực đầy đủ các chế độ retrieval và answering. Repo hiện tại hỗ trợ hybrid retrieval, contextual_hybrid retrieval, hybrid_rag, contextual_hybrid_rag và llm_only. Nhờ đó, cùng một câu hỏi có thể được đánh giá ở nhiều mức: chất lượng truy hồi, chất lượng grounded answer, và chênh lệch so với closed-book answer.",
            ),
            (
                "Normal",
                "Mục tiêu thứ ba là xây dựng bộ đo lường định lượng. Ở lớp retrieval, đề tài dùng recall@10, MRR@10, nDCG@10 và MAP. Ở lớp answer generation, hệ thống dùng faithfulness, correctness, relevancy và pairwise comparison. Cách tiếp cận này giúp kết luận không chỉ dựa vào cảm giác đọc câu trả lời mà còn dựa vào artifact có thể kiểm tra lại.",
            ),
            (
                "Normal",
                "Mục tiêu thứ tư là hoàn thiện một demo app cục bộ để phục vụ minh họa. Giao diện cho phép nhập truy vấn, xem top evidence, citation, kết quả từng mode, phase overview và failure cases. Nhờ đó, báo cáo không tách rời khỏi repo, mà phản ánh đúng những gì hệ thống hiện đang làm được.",
            ),
        ],
        "Giới hạn đề tài": [
            (
                "Normal",
                "Giới hạn thứ nhất nằm ở phạm vi bài toán. Đề tài chỉ tập trung vào nutrition-related health information, tức các câu hỏi sức khỏe phổ thông có liên hệ đến chế độ ăn, thực phẩm, vi chất, nguy cơ bệnh và bằng chứng từ tài liệu công khai. Hệ thống không được thiết kế cho chẩn đoán lâm sàng, điều trị cá thể hóa hay tư vấn y khoa chuyên sâu.",
            ),
            (
                "Normal",
                "Giới hạn thứ hai nằm ở dữ liệu. Corpus hiện tại gồm 2273 documents và 2320 nodes, trong đó NFCorpus chiếm phần lớn. Cách xây dựng này phù hợp cho benchmark và demo cục bộ, nhưng chưa thể xem là kho tri thức y sinh toàn diện. Một số chủ đề có coverage tốt, trong khi một số chủ đề vẫn còn ít evidence hoặc thiếu nguồn authoritative đa dạng.",
            ),
            (
                "Normal",
                "Giới hạn thứ ba nằm ở hạ tầng chạy thực nghiệm. Pha retrieval và answer generation phụ thuộc vào dịch vụ mô hình ngoài, còn answer-level evaluation và pairwise comparison có thể bị ảnh hưởng bởi quota API. Vì vậy, trong báo cáo này, những phần có artifact đầy đủ sẽ được phân tích định lượng, còn những phần chưa có đủ tín hiệu sẽ được trình bày theo đúng trạng thái thực tế thay vì suy diễn quá mức.",
            ),
        ],
        "Bố cục báo cáo": [
            ("Normal", "Báo cáo được tổ chức thành năm chương."),
            ("List Paragraph", "Chương 1 giới thiệu bài toán, mục tiêu, phạm vi và giới hạn của đề tài."),
            (
                "List Paragraph",
                "Chương 2 trình bày cơ sở lý thuyết liên quan đến truy hồi thông tin, vector retrieval, hybrid retrieval, RAG và các công trình nghiên cứu có liên quan đến medical QA.",
            ),
            (
                "List Paragraph",
                "Chương 3 mô tả hệ thống đã được triển khai trong repo, bao gồm pipeline bốn pha, knowledge base, schema dữ liệu, chiến lược truy hồi và demo app.",
            ),
            (
                "List Paragraph",
                "Chương 4 trình bày thiết lập thực nghiệm, phương pháp đánh giá, kết quả retrieval, kết quả answer-level evaluation, ví dụ minh họa và phần thảo luận.",
            ),
            ("List Paragraph", "Chương 5 tổng kết kết quả chính, nêu ưu điểm, hạn chế và các hướng phát triển tiếp theo."),
        ],
    }
)

SECTION_CONTENT.update(
    {
        "Thiết lập các cài đặt mô hình": [
            (
                "Normal",
                "Cấu hình mẫu của hệ thống được lưu trong .env.example. Ở thời điểm viết báo cáo, embedding model mặc định là text-embedding-3-small, còn mô hình dùng cho generation và evaluation là gpt-4.1-mini. Contextual retrieval được bật mặc định với target chunk khoảng 700 tokens, overlap 100 tokens, context max output 120 tokens và contextual candidate-k là 150. Reranker Cohere đã được cấu hình nhưng để ở trạng thái tắt mặc định.",
            ),
            (
                "Normal",
                "Bộ thực nghiệm được dùng để cập nhật báo cáo lần này được tổng hợp trong file output/contextual_retrieval_results_2026-04-05.md. Run này chấm trên 307 benchmark queries, dùng batch retrieval chính 7d8a1026-3f1f-483a-8e31-58e2ddaa1767 và bao gồm đủ các mode hybrid, contextual_hybrid, hybrid_rag, contextual_hybrid_rag và llm_only. Nhờ đó, Chương 4 có thể dựa trên full run thay vì chỉ dựa trên snapshot nhỏ hoặc các artifact chưa hoàn chỉnh như trước.",
            ),
            (
                "Normal",
                "Ở lớp demo, hệ thống thường hiển thị số citation ngắn gọn hơn để thuận tiện cho người dùng đọc nhanh, trong khi lớp benchmark dùng top-k = 10 để có cơ sở tính recall, MRR, nDCG và MAP. Vì vậy, phần demo app trong báo cáo chỉ đóng vai trò minh họa trực quan, còn kết luận định lượng chủ yếu dựa vào artifact đánh giá batch đầy đủ.",
            ),
        ],
        "Đánh giá 2 phương pháp": [
            (
                "Normal",
                "Đề tài đánh giá hai hướng tiếp cận chính là grounded answering và non-grounded answering. Trong grounded branch, repo hiện có hai biến thể là hybrid_rag và contextual_hybrid_rag. Trong non-grounded branch, hệ thống dùng llm_only làm baseline closed-book. Cách chia này giúp so sánh bản chất của hai hướng trả lời, đồng thời vẫn giữ được khả năng phân tích sâu hơn cho grounded branch.",
            ),
            (
                "Normal",
                "Lớp đánh giá thứ nhất là retrieval evaluation, áp dụng cho mode hybrid và contextual_hybrid. Bốn chỉ số được dùng là recall@10, MRR@10, nDCG@10 và MAP. Với contextual_hybrid, hệ thống còn bổ sung unique_docs@10 và duplicate_chunks@10 để theo dõi xem contextual retrieval có làm tăng tình trạng trùng lặp chunk trong top-k hay không.",
            ),
            (
                "Normal",
                "Lớp đánh giá thứ hai là answer-level evaluation. Ở lớp này, mỗi câu trả lời được chấm theo ba chiều: faithfulness, correctness và relevancy. Faithfulness đo mức độ bám bằng chứng truy hồi; correctness đo mức độ đúng của nội dung trả lời; còn relevancy đo mức độ trả lời đúng trọng tâm câu hỏi. Đây là ba góc nhìn bổ sung cho nhau, vì một câu trả lời có thể đúng một phần nhưng lan man, hoặc liên quan nhưng chưa đủ căn cứ.",
            ),
            (
                "Normal",
                "Lớp đánh giá thứ ba là pairwise comparison, trong đó evaluator so sánh trực tiếp hai câu trả lời cho cùng một truy vấn. Về lý thuyết, cách đo này rất hữu ích khi cần biết người chấm thích câu trả lời nào hơn trong bối cảnh thực tế. Tuy nhiên, chất lượng của lớp đánh giá này phụ thuộc vào việc artifact pairwise có được sinh đầy đủ hay không.",
            ),
            (
                "Normal",
                "Một lưu ý phương pháp quan trọng là llm_only không có retrieval context, nên faithfulness theo nghĩa grounded không phản ánh đúng bản chất của mode này. Vì vậy, các con số faithfulness = 0.0 ở llm_only cần được hiểu như trạng thái artifact của evaluator, không phải bằng chứng cho thấy câu trả lời closed-book luôn sai hoàn toàn. Khi phân tích, báo cáo sẽ tách rõ điều này để tránh diễn giải nhầm.",
            ),
        ],
        "4.3.1. Kết quả retrieval": [
            ("Caption", "Bảng 4.1. Kết quả retrieval trên batch đánh giá 307 truy vấn."),
            (
                "Normal",
                "Kết quả retrieval cho thấy hybrid và contextual_hybrid có chất lượng khá gần nhau trên batch 307 truy vấn. Hybrid đạt recall@10 = 0.195429, MRR@10 = 0.548713, nDCG@10 = 0.353326 và MAP = 0.158659. Contextual_hybrid đạt recall@10 = 0.193321, MRR@10 = 0.531045, nDCG@10 = 0.341681 và MAP = 0.155346.",
            ),
            (
                "Normal",
                "Nhìn theo trung bình batch, hybrid hiện nhỉnh hơn contextual_hybrid ở các retrieval metrics chính. Tuy vậy, contextual_hybrid vẫn cho thấy đặc điểm đáng chú ý là unique_docs@10 trung bình đạt 9.9023 và duplicate_chunks@10 chỉ 0.0977, tức top-k vẫn khá đa dạng tài liệu. Điều này gợi ý rằng contextual retrieval không phá vỡ tính đa dạng của candidate, nhưng lợi ích định lượng của nó còn phụ thuộc vào từng truy vấn cụ thể.",
            ),
        ],
        "Kết quả answer-level evaluation": [
            ("Caption", "Bảng 4.2. Kết quả answer-level evaluation của các chế độ trả lời."),
            (
                "Normal",
                "Ở answer-level evaluation, hybrid_rag là mode có artifact đầy đủ nhất. Trên 307 truy vấn, mode này đạt faithfulness trung bình 0.9548, correctness 0.8224 và relevancy 0.8876. Ba chỉ số này cho thấy grounded answer thường bám evidence tốt và trả lời khá đúng trọng tâm trong phạm vi corpus hiện tại.",
            ),
            (
                "Normal",
                "Ở full run ngày 2026-04-05, contextual_hybrid_rag đạt faithfulness 0.9493, correctness 0.8284 và relevancy 0.9216 trên 307 truy vấn. So với hybrid_rag, contextual_hybrid_rag giảm nhẹ ở faithfulness nhưng tăng ở correctness và relevancy. Điều này cho thấy contextual retrieval chưa giúp hệ thống kéo được nhiều tài liệu đúng hơn ở lớp ranking, nhưng lại giúp tầng generation trả lời trúng trọng tâm hơn trong nhiều trường hợp.",
            ),
            (
                "Normal",
                "Đối với llm_only, full run hiện đã có đủ 307 dòng answer-level evaluation với correctness trung bình 0.4243 và relevancy 0.9911, trong khi faithfulness vẫn bằng 0.0000 theo đúng bản chất của metric grounded. Cách đọc phù hợp ở đây là llm_only vẫn có thể trả lời rất sát trọng tâm câu hỏi về mặt diễn đạt, nhưng khi đối chiếu với evidence benchmark thì độ đúng nội dung còn thấp hơn đáng kể so với hai mode grounded.",
            ),
            ("Caption", "Bảng 4.3. Tóm tắt pairwise comparison giữa các chế độ trả lời."),
        ],
        "4.3.3. Pairwise comparison": [
            (
                "Normal",
                "Khác với trạng thái trước đó, full run ngày 2026-04-05 đã cho ra hai kết quả pairwise tương đối sạch. Hybrid_rag thắng llm_only ở 223 trên 307 truy vấn, tương ứng left-win rate 0.7264. Mạnh hơn nữa, contextual_hybrid_rag thắng hybrid_rag ở 276 trên 307 truy vấn, tương ứng left-win rate 0.8990. Đây là tín hiệu định lượng mạnh nhất của run hiện tại và cũng là bằng chứng rõ nhất cho lợi ích của contextual retrieval ở tầng answer generation.",
            ),
            (
                "Normal",
                "Cặp contextual_hybrid_rag so với llm_only cũng đã có kết quả sơ bộ là 169 trên 307 truy vấn, tương ứng left-win rate 0.5505. Tuy nhiên, file tổng hợp cũng ghi rõ còn 72 dòng bị ảnh hưởng bởi quota-related judge failures, nên cặp so sánh này chỉ nên được xem là tín hiệu tham khảo. Vì vậy, trong phân tích, báo cáo sẽ xem cặp contextual_hybrid_rag vs hybrid_rag và hybrid_rag vs llm_only là hai kết quả pairwise đáng tin cậy hơn.",
            ),
        ],
        "4.3.4. Truy vấn minh họa trên demo app": [
            (
                "Normal",
                "Ở truy vấn minh họa PLAIN-2, câu hỏi “Do Cholesterol Statin Drugs Cause Breast Cancer?” cho thấy rõ vai trò của retrieval. Top evidence của hybrid và contextual_hybrid đều xoay quanh plant sterols, statin, cholesterol metabolism và nguy cơ ung thư vú. Grounded answer từ hai mode retrieval đều kết luận rằng evidence hiện có không ủng hộ mệnh đề statin gây ung thư vú, đồng thời giữ giọng điệu thận trọng vì các tài liệu liên quan trực tiếp đến statin và breast cancer không nhiều.",
            ),
            (
                "Normal",
                "Trong khi đó, llm_only cũng trả lời rằng statin chưa được chứng minh là gây ung thư vú, nhưng không đưa ra citation hay evidence bundle đi kèm. Ví dụ này minh họa khá rõ giá trị thực tế của retrieval trong hệ thống: không phải lúc nào grounded answer cũng khác hoàn toàn về kết luận, nhưng nó minh bạch hơn ở chỗ cho phép người đọc kiểm tra câu trả lời đang dựa trên tài liệu nào.",
            ),
        ],
        "Phân tích và thảo luận kết quả": [
            (
                "Normal",
                "Kết quả retrieval cho thấy bài toán vẫn còn khó ở khía cạnh coverage. Dù MRR@10 của hybrid đạt 0.548713, recall@10 và MAP chỉ ở mức khoảng 0.19 và 0.16. Điều này phản ánh một hiện tượng quen thuộc trong IR: hệ thống thường tìm được một vài tài liệu rất gần truy vấn, nhưng chưa thu hồi đầy đủ toàn bộ tài liệu liên quan trong top-k.",
            ),
            (
                "Normal",
                "Contextual_hybrid không vượt hybrid ở retrieval metrics chính: recall@10, MRR@10, nDCG@10 và MAP đều thấp hơn nhẹ. Điều này có nghĩa là nếu chỉ nhìn ở lớp ranking thuần túy, contextual retrieval chưa tạo ra bước nhảy rõ rệt. Tuy nhiên, answer-level evaluation lại cho thấy contextual_hybrid_rag tăng correctness từ 0.8224 lên 0.8284 và tăng relevancy từ 0.8876 lên 0.9216. Đây là một kết quả đáng chú ý vì nó cho thấy contextualization có thể giúp mô hình sử dụng evidence tốt hơn, ngay cả khi thứ hạng retrieval chưa được cải thiện.",
            ),
            (
                "Normal",
                "Kết quả pairwise củng cố nhận định trên. Việc contextual_hybrid_rag thắng hybrid_rag với left-win rate 0.8990 cho thấy contextual retrieval không chỉ cải thiện một vài metric đơn lẻ, mà còn làm evaluator thường xuyên ưu tiên câu trả lời contextualized khi so sánh trực tiếp. Đồng thời, hybrid_rag thắng llm_only với left-win rate 0.7264, cho thấy grounded answering nhìn chung vẫn đáng tin hơn closed-book answering trên benchmark hiện tại.",
            ),
            (
                "Normal",
                "Một điểm tinh tế hơn là llm_only có relevancy rất cao, đạt 0.9911, nhưng correctness chỉ 0.4243. Điều này phản ánh một hiện tượng quen thuộc của LLM-only: mô hình có thể bám câu hỏi rất tốt về mặt ngôn ngữ và cách diễn đạt, nhưng vẫn trả lời sai hoặc lệch so với benchmark evidence. Trong khi đó, grounded modes có thể diễn đạt thận trọng hơn hoặc dài hơn, nhưng lại đúng hơn ở góc nhìn benchmark. Vì vậy, nếu chỉ nhìn câu chữ bề ngoài thì llm_only dễ gây cảm giác tốt, còn nếu nhìn bằng evidence-based evaluation thì grounded branch rõ ràng mạnh hơn.",
            ),
            (
                "Normal",
                "Các failure cases trong file tổng hợp cũng cho thấy hướng cải thiện khá rõ. Ở retrieval, hybrid còn hụt ở các truy vấn như Best Treatment for Constipation hoặc Increasing Muscle Strength with Fenugreek, còn contextual_hybrid hụt ở các truy vấn như Chronic Headaches and Pork Parasites hoặc Plant-Based Diets for Psoriasis. Ở answer generation, hybrid_rag còn có xu hướng bỏ sót evidence ở một số truy vấn như St. John's wort, trong khi contextual_hybrid_rag vẫn có thể thất bại ở các truy vấn như rickets hoặc breast pain. Điều này cho thấy contextual retrieval là một cải tiến hữu ích, nhưng chưa giải quyết triệt để toàn bộ failure modes.",
            ),
        ],
        "Nhận xét kết quả đề tài": [
            (
                "Normal",
                "Đề tài đã xây dựng thành công một pipeline tương đối đầy đủ cho bài toán nutrition-focused health QA, bao gồm chuẩn bị dữ liệu, indexing, retrieval, answer generation, evaluation và demo app. So với một hệ thống chỉ dừng ở mức proof-of-concept, repo hiện tại có ưu điểm là mỗi pha đều sinh ra artifact kiểm tra được, từ đó giúp báo cáo có cơ sở mô tả và đối chiếu với mã nguồn thực tế.",
            ),
            (
                "Normal",
                "Về mặt kết quả, full run hiện tại cho thấy một kết luận cân bằng hơn. Contextual retrieval chưa làm retrieval metrics tốt hơn hybrid baseline, nhưng contextual_hybrid_rag lại cải thiện correctness, relevancy và đặc biệt là pairwise preference so với hybrid_rag. Đồng thời, grounded answering vẫn thắng llm_only khá rõ ở pairwise comparison. Vì vậy, có thể kết luận rằng giá trị lớn nhất của contextual retrieval trong repo hiện tại xuất hiện ở tầng dùng evidence để sinh câu trả lời, hơn là ở tầng xếp hạng thuần túy.",
            ),
        ],
        "Ưu – nhược điểm của đề tài": [
            (
                "Normal",
                "Ưu điểm lớn nhất của đề tài là tính tái hiện. Repo được chia pha rõ ràng, có schema PostgreSQL nhất quán, có runtime log, có metric và có demo app. Bên cạnh đó, việc đặt BM25, vector search, contextual retrieval và experiment tracking trong cùng một hệ thống giúp toàn bộ pipeline dễ kiểm chứng hơn so với cách ghép nhiều dịch vụ rời rạc.",
            ),
            (
                "Normal",
                "Ưu điểm thứ hai là tính minh bạch của grounded answer. So với llm_only, các mode grounded cho phép quan sát evidence bundle và citation, nhờ đó người đọc có thể đánh giá câu trả lời ở cả tầng nội dung lẫn tầng nguồn chứng cứ. Đây là điểm đặc biệt quan trọng trong miền sức khỏe, nơi một câu trả lời đúng về nội dung nhưng không có khả năng truy vết vẫn chưa phải lời giải lý tưởng.",
            ),
            (
                "Normal",
                "Nhược điểm chính của đề tài hiện không còn nằm ở việc thiếu kết quả, mà nằm ở việc kết quả cho thấy trade-off rõ ràng. Contextual branch cải thiện answer quality nhưng chưa cải thiện retrieval metrics. Cặp contextual_hybrid_rag vs llm_only vẫn còn 72 dòng bị quota-related judge failures, nên chưa thể xem là kết luận cuối cùng. Ngoài ra, retrieval metrics cho thấy hệ thống vẫn còn bỏ sót nhiều tài liệu liên quan trong top-k, nghĩa là knowledge base, chunking và lexical-semantic fusion vẫn còn không gian để tối ưu thêm.",
            ),
        ],
        "Hướng phát triển": [
            (
                "Normal",
                "Hướng phát triển đầu tiên là làm sạch và mở rộng lớp thực nghiệm. Ưu tiên gần nhất là xử lý nốt 72 dòng quota-affected của cặp contextual_hybrid_rag vs llm_only để có một phép so sánh pairwise hoàn chỉnh. Sau đó, nhóm có thể bổ sung thêm batch kiểm thử mới hoặc chia benchmark theo nhóm truy vấn để xem contextual retrieval mạnh ở loại câu hỏi nào thay vì chỉ nhìn trung bình toàn cục.",
            ),
            (
                "Normal",
                "Hướng thứ hai là tối ưu retrieval theo failure cases. Vì contextual retrieval hiện mạnh hơn ở lớp answer hơn là lớp ranking, nhóm có thể thử tinh chỉnh lại chunk target, overlap, context summary prompt và candidate-k để xem retrieval metrics có thể cải thiện hay không. Song song với đó, việc bật reranker trong các cấu hình có kiểm soát, mở rộng nguồn authoritative và thử iterative retrieval hoặc query rewriting cũng là những bước hợp lý để tăng coverage cho các truy vấn khó.",
            ),
        ],
    }
)

REFERENCE_LINES = [
    "[1] S. Robertson and H. Zaragoza, “The probabilistic relevance framework: BM25 and beyond,” Foundations and Trends in Information Retrieval, vol. 3, no. 4, pp. 333-389, 2009, doi: 10.1561/1500000019.",
    "[2] V. Karpukhin, B. Oğuz, S. Min, P. Lewis, L. Wu, S. Edunov, D. Chen, and W.-t. Yih, “Dense passage retrieval for open-domain question answering,” in Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing (EMNLP), 2020, doi: 10.48550/arXiv.2004.04906.",
    "[3] P. Lewis, E. Perez, A. Piktus, F. Petroni, V. Karpukhin, N. Goyal, H. Küttler, M. Lewis, W.-t. Yih, T. Rocktäschel, S. Riedel, and D. Kiela, “Retrieval-augmented generation for knowledge-intensive NLP tasks,” in Advances in Neural Information Processing Systems (NeurIPS), 2020, doi: 10.48550/arXiv.2005.11401.",
    "[4] K. Singhal et al., “Large language models encode clinical knowledge,” Nature, vol. 620, pp. 172-180, 2023, doi: 10.1038/s41586-023-06291-2.",
    "[5] H. Nori, N. King, S. M. McKinney, D. Carignan, and E. Horvitz, “Capabilities of GPT-4 on medical challenge problems,” arXiv preprint, 2023, doi: 10.48550/arXiv.2303.13375.",
    "[6] Y. Cai et al., “MedBench: A large-scale Chinese benchmark for evaluating medical large language models,” arXiv:2312.12806, 2023. [Online]. Available: https://arxiv.org/abs/2312.12806. [Accessed: Apr. 5, 2026].",
    "[7] G. Xiong, Q. Jin, Z. Lu, and A. Zhang, “Benchmarking retrieval-augmented generation for medicine,” arXiv:2402.13178, 2024. [Online]. Available: https://arxiv.org/abs/2402.13178. [Accessed: Apr. 5, 2026].",
    "[8] I. Alonso, M. Oronoz, and R. Agerri, “MedExpQA: Multilingual benchmarking of large language models for medical question answering,” arXiv:2404.05590, 2024. [Online]. Available: https://arxiv.org/abs/2404.05590. [Accessed: Apr. 5, 2026].",
    "[9] Y. Kim, J. Wu, Y. Abdulle, and H. Wu, “MedExQA: Medical question answering benchmark with multiple explanations,” arXiv:2406.06331, 2024. [Online]. Available: https://arxiv.org/abs/2406.06331. [Accessed: Apr. 5, 2026].",
    "[10] G. Xiong, Q. Jin, X. Wang, M. Zhang, Z. Lu, and A. Zhang, “Improving retrieval-augmented generation in medicine with iterative follow-up questions,” arXiv:2408.00727, 2024. [Online]. Available: https://arxiv.org/abs/2408.00727. [Accessed: Apr. 5, 2026].",
    "[11] S. Liang, L. Zhang, H. Zhu, W. Wang, Y. He, and D. Zhou, “RGAR: Recurrence generation-augmented retrieval for factual-aware medical question answering,” arXiv:2502.13361, 2025. [Online]. Available: https://arxiv.org/abs/2502.13361. [Accessed: Apr. 5, 2026].",
    "[12] D. Ford, “Introducing Contextual Retrieval,” Anthropic Engineering, Sep. 19, 2024. [Online]. Available: https://www.anthropic.com/engineering/contextual-retrieval. [Accessed: Apr. 5, 2026].",
    "[13] A. Sewrathan, “PostgreSQL as a Vector Database: A pgvector Tutorial,” Tiger Data Blog, Jun. 21, 2023. [Online]. Available: https://www.tigerdata.com/blog/postgresql-as-a-vector-database-using-pgvector. [Accessed: Apr. 5, 2026].",
    "[14] Tiger Data Documentation, “Optimize full text search with BM25.” [Online]. Available: https://www.tigerdata.com/docs/use-timescale/latest/extensions/pg-textsearch. [Accessed: Apr. 5, 2026].",
]

SECTION_CONTENT.update(
    {
        "Pipeline tổng thể": [
            ("List Paragraph", "1. Thu thập và chuẩn hóa dữ liệu từ NFCorpus, MedlinePlus, FDA và PubMed thành ba file nguồn JSONL."),
            ("List Paragraph", "2. Chuẩn hóa dữ liệu thành hai lớp chính là documents và nodes để phục vụ retrieval."),
            ("List Paragraph", "3. Sinh contextual summary cho từng chunk, tạo body thường và contextualized_body, rồi ghi dữ liệu vào PostgreSQL."),
            ("List Paragraph", "4. Tạo bốn chỉ mục chính cho retrieval: BM25, contextual BM25, vector embedding và contextual embedding."),
            ("List Paragraph", "5. Nhận câu hỏi đầu vào, chạy hybrid retrieval hoặc contextual_hybrid retrieval, sau đó fuse bằng Reciprocal Rank Fusion."),
            ("List Paragraph", "6. Sinh câu trả lời theo các chế độ hybrid_rag, contextual_hybrid_rag hoặc llm_only, đồng thời lưu evidence bundle và citations."),
            ("List Paragraph", "7. Đánh giá retrieval, answer quality và pairwise comparison, sau đó tổng hợp system summary để phục vụ dashboard."),
            ("Style1", "Mô tả từng module"),
            ("Normal", "Module 01 - Data Preparation."),
            (
                "Normal",
                "Module này tạo ra corpora đầu vào phục vụ toàn bộ hệ thống. Dữ liệu benchmark từ NFCorpus được lọc theo chủ đề dinh dưỡng, dữ liệu từ MedlinePlus và FDA được crawl từ nguồn chính thống, còn dữ liệu PubMed được chuẩn hóa lại để đưa về cùng schema. Kết quả đầu ra là ba file nguồn gồm nfcorpus_nutrition.jsonl, nutrition_crawl.jsonl và pubmed_kb.jsonl.",
            ),
            ("Normal", "Module 02 - Indexing."),
            (
                "Normal",
                "Module indexing đọc ba file JSONL đầu vào, tạo documents và nodes, ước lượng token, chia chunk có overlap khi cần, sinh contextual summary, export benchmark queries và qrels, sau đó upsert dữ liệu vào các bảng kb_documents và kb_nodes trong PostgreSQL. Ở bước này, hệ thống cũng chuẩn bị cả embedding thường và contextual embedding để phục vụ hai nhánh retrieval khác nhau.",
            ),
            (
                "Normal",
                "Về mặt cài đặt, phase 2 có các lệnh bootstrap, prepare-records, export-benchmark, index-postgres và index-vector. Các tham số cấu hình quan trọng của phase này nằm trong file .env, đặc biệt là CONTEXTUAL_RETRIEVAL_ENABLED = true, CONTEXTUAL_RETRIEVAL_LLM_MODEL = gpt-4.1-mini, CONTEXTUAL_CHUNK_TARGET_TOKENS = 700, CONTEXTUAL_CHUNK_OVERLAP_TOKENS = 100, CONTEXTUAL_CONTEXT_MAX_TOKENS = 120 và CONTEXTUAL_DOCUMENT_CHAR_LIMIT = 16000. Đây là nhóm biến quyết định cách hệ thống chia chunk, lấy ngữ cảnh tài liệu và sinh context summary trước khi index.",
            ),
            ("Normal", "Module 03 - Retrieval."),
            (
                "Normal",
                "Module retrieval hiện hỗ trợ các mode hybrid, contextual_hybrid, hybrid_rag, contextual_hybrid_rag và llm_only. Mỗi lần chạy đều ghi runtime log ra bảng PostgreSQL và file JSONL, nhờ đó kết quả có thể tái sử dụng cho phần đánh giá mà không cần tái hiện toàn bộ pipeline.",
            ),
            (
                "Normal",
                "Phase 3 có các lệnh chính bootstrap, db-init, ingest-postgres, query, batch, clear-results và status. Ở lớp retrieval, các tham số quan trọng gồm CONTEXTUAL_RETRIEVAL_CANDIDATE_K = 150 để quyết định số candidate được lấy ở contextual_hybrid, top-k trong lệnh query hoặc batch để quyết định số bằng chứng cuối cùng, cùng nhóm RERANKER_ENABLED, RERANKER_PROVIDER, RERANKER_CANDIDATE_K, COHERE_API_KEY, COHERE_BASE_URL và COHERE_RERANK_MODEL nếu bật reranker. Trong cấu hình mẫu hiện tại, RERANKER_ENABLED = false, RERANKER_PROVIDER = cohere, COHERE_RERANK_MODEL = rerank-v4.0-fast và RERANKER_CANDIDATE_K = 20.",
            ),
            ("Normal", "Module 04 - Evaluation."),
            (
                "Normal",
                "Module evaluation đọc retrieval_runs, answer_runs, answer_evaluations và comparison_runs để tính metric, tổng hợp failure cases và sinh system_summary.json. Đây là lớp biến raw runtime log thành tín hiệu có thể đưa vào báo cáo và demo dashboard.",
            ),
            (
                "Normal",
                "Phase 4 có các lệnh retrieval-metrics, answer-eval, compare và report. Pha này không sinh thêm retrieval context mới, nhưng vẫn phụ thuộc vào OPENAI_API_KEY, OPENAI_BASE_URL và OPENAI_LLM_MODEL để gọi evaluator của DeepEval. Vì vậy, nếu phase 3 dùng gpt-4.1-mini làm generation model thì phase 4 mặc định cũng dùng gpt-4.1-mini làm evaluation model, trừ khi người dùng đổi biến môi trường tương ứng.",
            ),
            ("Normal", "Module 05 - Demo App."),
            (
                "Normal",
                "Lớp demo app gồm FastAPI backend và Next.js frontend. Backend bọc runtime Python thành các API như /health, /kb/summary, /demo/query, /demo/summary, /demo/phases và /demo/failure-cases. Frontend dùng các API này để hiển thị dashboard, phase overview, workbench truy vấn, evidence bundle, citations và failure cases theo dạng trực quan hơn so với việc đọc file JSONL thủ công.",
            ),
            (
                "Normal",
                "Về cấu hình vận hành, run_demo.ps1 cho phép đặt ApiPort và WebPort, đồng thời gán MEDIR_DEMO_API_URL để frontend gọi đúng backend cục bộ. Lớp demo không thêm tham số retrieval mới, mà đọc lại các cấu hình và artifact đã được phase 2, phase 3 và phase 4 tạo sẵn. Điều này giúp giao diện demo phản ánh đúng cấu hình đã chạy, thay vì dùng một bộ tham số riêng không xuất hiện trong pipeline chính.",
            ),
        ],
        "Công cụ hiện thực": [
            (
                "Normal",
                "Hệ thống sử dụng kết hợp thư viện xử lý dữ liệu, cơ sở dữ liệu, hạ tầng API và web frontend. Bảng 3.1 tóm tắt các thành phần chính đang được dùng trong repo hiện tại.",
            ),
            ("Caption", "Bảng 3.1. Thành phần công nghệ sử dụng trong hệ thống."),
            (
                "Normal",
                "Trong cấu hình mẫu hiện tại, embedding model là text-embedding-3-small, còn mô hình sinh câu trả lời và evaluator dùng gpt-4.1-mini. PostgreSQL đóng vai trò vừa là nơi lưu knowledge base, vừa là nơi lưu runtime logs, còn pgvector và pg_textsearch giúp lexical retrieval và semantic retrieval cùng nằm trong một hạ tầng thống nhất [13], [14]. Cohere reranker đã được chuẩn bị ở mức tùy chọn để mở rộng thí nghiệm, nhưng mặc định trong .env.example tính năng này đang tắt.",
            ),
            (
                "Normal",
                "Nếu nhìn theo nhóm cấu hình, hệ thống có bốn lớp biến chính. Lớp database gồm POSTGRES_HOST, POSTGRES_PORT, POSTGRES_DB, POSTGRES_USER và POSTGRES_PASSWORD. Lớp model gồm OPENAI_API_KEY, OPENAI_BASE_URL, OPENAI_EMBED_MODEL và OPENAI_LLM_MODEL. Lớp contextual retrieval gồm CONTEXTUAL_RETRIEVAL_ENABLED, CONTEXTUAL_RETRIEVAL_LLM_MODEL, CONTEXTUAL_CHUNK_TARGET_TOKENS, CONTEXTUAL_CHUNK_OVERLAP_TOKENS, CONTEXTUAL_CONTEXT_MAX_TOKENS, CONTEXTUAL_DOCUMENT_CHAR_LIMIT và CONTEXTUAL_RETRIEVAL_CANDIDATE_K. Lớp reranker gồm RERANKER_ENABLED, RERANKER_PROVIDER, RERANKER_CANDIDATE_K, COHERE_API_KEY, COHERE_BASE_URL và COHERE_RERANK_MODEL.",
            ),
            (
                "Normal",
                "Các biến trên không xuất hiện rời rạc mà đi theo phase. Phase 2 dùng mạnh nhóm contextual để quyết định cách chunking và contextualization. Phase 3 dùng đồng thời nhóm model, contextual và reranker để điều khiển truy hồi, fusion và answer generation. Phase 4 tái sử dụng nhóm model để gọi evaluator. Nhờ cách chia này, cùng một file .env có thể điều khiển toàn bộ pipeline mà vẫn dễ truy vết biến nào tác động lên bước nào.",
            ),
        ],
        "Tập dữ liệu": [
            (
                "Normal",
                "Sau bước chuẩn bị và hợp nhất dữ liệu, knowledge base hiện có 2273 documents và 2320 nodes. Chênh lệch giữa hai con số cho thấy một phần tài liệu đã được tách thành nhiều chunk; thống kê hiện tại cho thấy có 35 documents có chunk_count lớn hơn 1, và tài liệu dài nhất được tách thành 5 chunk.",
            ),
            ("Caption", "Bảng 3.2. Thống kê các nguồn dữ liệu được index."),
            (
                "Normal",
                "Ngoài knowledge base, pha indexing còn export 307 benchmark queries và 7773 dòng qrels test. Đây là nền tảng để tính retrieval metrics trên batch đủ lớn, thay vì chỉ dựa vào một vài truy vấn minh họa. Vì benchmark backbone đến từ NFCorpus, kết quả retrieval mang tính so sánh trong phạm vi corpus hiện tại hơn là một kết luận tuyệt đối cho toàn bộ miền y sinh.",
            ),
            (
                "Normal",
                "Một điểm cần nhấn mạnh là corpus của đề tài có tính mixed-purpose. NFCorpus cung cấp backbone benchmark và độ phủ truy vấn, MedlinePlus cùng FDA bổ sung nguồn public-health có tính chính thống, còn PubMed đóng vai trò đưa thêm góc nhìn nghiên cứu. Cách phối hợp này giúp hệ thống vừa có thể benchmark, vừa có evidence đủ đa dạng cho grounded answer, dù vẫn còn giới hạn về độ rộng chủ đề.",
            ),
        ],
        "Quy trình xây dựng Knowledge Base trong đề tài": [
            (
                "Normal",
                "Quy trình xây knowledge base của đề tài bắt đầu từ việc xác định use-case là nutrition-related health information. Từ use-case này, nhóm chọn nguồn theo ba lớp: benchmark corpus để có qrels và truy vấn chuẩn, nguồn public-health để tăng độ tin cậy, và bài báo nghiên cứu để mở rộng evidence cho các câu hỏi cần nhiều ngữ cảnh hơn. Sau khi thu thập, dữ liệu được chuẩn hóa về schema chung và đưa vào phase indexing.",
            ),
            (
                "Normal",
                "Ở mức vật lý, knowledge base hiện được chia thành hai bảng chính là kb_documents và kb_nodes. Bảng kb_documents lưu thông tin ở mức tài liệu gốc như doc_id, source_id, title, source_url, document_text, document_token_count, chunk_count và section_type. Bảng kb_nodes lưu đơn vị truy hồi thực tế như node_id, doc_id, body, raw_body, context_summary, contextualized_body, chunk_index, chunk_count, token_count, char_count, embedding và contextual_embedding.",
            ),
            (
                "Normal",
                "Thiết kế này cho phép một tài liệu vừa được xem như một record toàn cục, vừa được chia nhỏ thành các node phục vụ retrieval. Trường raw_body giữ văn bản chunk nguyên gốc, còn contextualized_body ghép thêm context summary để hỗ trợ contextual retrieval. Nhờ đó, cùng một node có thể được dùng cho hai chế độ truy hồi khác nhau mà không cần tạo hai knowledge base riêng biệt.",
            ),
            ("Caption", "Bảng 3.3. Thiết kế schema metadata cho Knowledge Base."),
            (
                "Normal",
                "Ngoài lớp dữ liệu nội dung, hệ thống còn có các bảng retrieval_runs, answer_runs, answer_evaluations và comparison_runs để lưu toàn bộ vòng đời thực nghiệm. Nhóm xem đây là một phần của knowledge base mở rộng theo nghĩa vận hành: không chỉ lưu tri thức, mà còn lưu dấu vết sử dụng tri thức để thuận tiện cho đánh giá, kiểm chứng và tái hiện kết quả.",
            ),
            (
                "Normal",
                "Về lifecycle, mỗi record đi qua các bước làm sạch, enrich metadata, chunking, contextualization, index, retrieval và evaluation. Cách làm này phản ánh khá rõ tinh thần của bài giảng về knowledge base trong IR: knowledge base chỉ thật sự hữu ích khi dữ liệu không chỉ được lưu, mà còn được tổ chức và quản trị theo cách phục vụ truy hồi hiệu quả.",
            ),
            (
                "Normal",
                "Ở pha cài đặt knowledge base, nhóm tham số ảnh hưởng trực tiếp nhất là nhóm contextual chunking. CONTEXTUAL_CHUNK_TARGET_TOKENS quyết định kích thước chunk mục tiêu; CONTEXTUAL_CHUNK_OVERLAP_TOKENS quyết định phần chồng lắp giữa hai chunk liên tiếp; CONTEXTUAL_CONTEXT_MAX_TOKENS quyết định độ dài tối đa của context summary; còn CONTEXTUAL_DOCUMENT_CHAR_LIMIT giới hạn phần văn bản tài liệu được đưa vào prompt contextualization. Vì vậy, khi báo cáo mô tả knowledge base, các tham số này cần được xem là một phần của thiết kế dữ liệu chứ không chỉ là chi tiết kỹ thuật phụ.",
            ),
            (
                "Normal",
                "Từ góc nhìn kỹ thuật, việc giữ metadata và vector trong cùng PostgreSQL giúp hệ thống đơn giản hóa hạ tầng. Thay vì tách lexical index, vector store và runtime tracking ra nhiều nơi khác nhau, repo hiện tại đặt chúng trong một cụm schema thống nhất. Đây là lựa chọn thực dụng, phù hợp với mục tiêu local-first và dễ tái hiện của đồ án [13], [14].",
            ),
        ],
        "Mô hình truy hồi và chiến lược trả lời": [
            (
                "Normal",
                "Ở tầng truy hồi, hệ thống dùng hai nguồn tín hiệu song song. Nhánh thứ nhất là BM25 chạy trên trường body của bảng kb_nodes thông qua pg_textsearch. Nhánh này đặc biệt hữu ích với các truy vấn chứa thực thể rõ ràng như tên chất, tên thuốc, tên bệnh hoặc cụm từ khóa đặc thù. Trong repo, truy vấn lexical được biểu diễn qua to_bm25query() và các chỉ mục bm25 được tạo trực tiếp trên PostgreSQL [14].",
            ),
            (
                "Normal",
                "Nếu đi sâu hơn vào lớp SQL, BM25 của hệ thống không phải một khái niệm trừu tượng mà được cài đặt trực tiếp bằng chỉ mục inverted index trên PostgreSQL. Trong file schema.sql, repo tạo hai chỉ mục gồm kb_nodes_bm25_idx trên cột body và kb_nodes_contextual_bm25_idx trên cột contextualized_body bằng cú pháp USING bm25 (...). Có thể hiểu đơn giản rằng ở bước index, văn bản của từng node được phân tách thành các term để tạo ra cấu trúc ánh xạ từ term sang danh sách tài liệu hoặc chunk có chứa term đó. Đây chính là bản chất của inverted index trong lexical retrieval.",
            ),
            (
                "Normal",
                "Quá trình phân tách term không diễn ra tùy ý mà chịu tác động của tham số text_config = 'english' được khai báo ngay khi tạo chỉ mục. Tham số này quyết định cách PostgreSQL chuẩn hóa văn bản trước khi lập chỉ mục, ví dụ cách tách token theo từ, chuẩn hóa một số biến thể từ và xử lý theo cấu hình ngôn ngữ tiếng Anh. Vì corpus của repo chủ yếu là tài liệu tiếng Anh, việc dùng text_config = 'english' là hợp lý: cùng một cơ chế tokenization sẽ được áp dụng nhất quán cho cả tài liệu lẫn truy vấn.",
            ),
            (
                "Normal",
                "Ở chiều truy vấn, retrieval.py dùng toán tử body <@> to_bm25query(%(query)s, 'kb_nodes_bm25_idx') cho BM25 thường và contextualized_body <@> to_bm25query(%(query)s, 'kb_nodes_contextual_bm25_idx') cho contextual BM25. Có thể xem to_bm25query() như bước chuyển câu hỏi thô của người dùng sang dạng BM25 query tương thích với inverted index đã tạo trước đó. Nhờ vậy, truy vấn và tài liệu được phân tích dưới cùng một cấu hình ngôn ngữ, giảm rủi ro lệch pha giữa cách token hóa lúc index và lúc search.",
            ),
            (
                "Normal",
                "Sự khác nhau giữa hai nhánh lexical nằm ở cột được lập chỉ mục. BM25 thường dùng body, tức văn bản chunk gần với nội dung gốc nhất. Contextual BM25 dùng contextualized_body, tức raw_body đã được ghép thêm context summary. Điều này làm cho cùng một chunk có thể mang thêm các token mô tả chủ đề rộng hơn của tài liệu, từ đó tăng cơ hội match với truy vấn ngay cả khi raw chunk ban đầu quá ngắn hoặc thiếu từ khóa bề mặt. Vì vậy, contextual retrieval trong repo không chỉ là thêm embedding contextual, mà còn là thêm một inverted index lexical mới trên văn bản đã được contextualize.",
            ),
            (
                "Normal",
                "Nhánh thứ hai là vector retrieval chạy trên cột embedding của kb_nodes. Embedding được sinh offline bằng text-embedding-3-small và được lưu dưới dạng vector(1536) trong PostgreSQL. Để tăng tốc tìm kiếm lân cận gần nhất, repo tạo HNSW index cho cả embedding thường và contextual_embedding. Cách làm này cho phép vector search và metadata filtering cùng diễn ra trong một hệ quản trị dữ liệu duy nhất [13].",
            ),
            (
                "Normal",
                "Sau khi có hai danh sách candidate, hệ thống dùng Reciprocal Rank Fusion để hợp nhất thứ hạng. RRF có ưu điểm là đơn giản, dễ giải thích và không cần học thêm tham số mới, nhưng vẫn tận dụng được thế mạnh bổ sung của lexical retrieval và semantic retrieval. Vì vậy, đây là lựa chọn phù hợp với repo thực nghiệm cần tính minh bạch cao.",
            ),
            (
                "Normal",
                "Phiên bản hiện tại còn bổ sung contextual retrieval. Ở bước indexing, mỗi chunk được sinh context_summary ngắn dựa trên tài liệu gốc. Sau đó, context_summary được ghép vào raw_body để tạo contextualized_body. Khi truy hồi, hệ thống có thể tìm kiếm trên contextualized_body và contextual_embedding, tức vừa tăng tín hiệu lexical, vừa tăng tín hiệu semantic cho những chunk vốn quá ngắn hoặc thiếu ngữ cảnh cục bộ [12].",
            ),
            (
                "Normal",
                "Ngoài ra, repo đã chuẩn bị sẵn một lớp reranker tùy chọn dùng Cohere. Nếu bật cấu hình tương ứng, candidate của từng nhánh sẽ được sắp xếp lại trước khi fuse. Tuy nhiên, do đây là tính năng optional và mặc định đang tắt trong .env.example, báo cáo chỉ mô tả nó như một phần mở rộng đã được chuẩn bị trong kiến trúc, không xem là cấu hình bắt buộc của mọi lần chạy.",
            ),
            (
                "Normal",
                "Nếu nhìn sát hơn vào code, contextual retrieval và reranker được điều khiển bằng các tham số khá rõ ràng. Contextual branch dùng CONTEXTUAL_RETRIEVAL_CANDIDATE_K = 150 để mở rộng số candidate quét trước khi chọn top-k cuối, giúp contextual BM25 và contextual embedding có không gian hợp nhất rộng hơn. Reranker, nếu bật, dùng RERANKER_CANDIDATE_K = 20 để quyết định số ứng viên ban đầu của mỗi nhánh được đem đi sắp lại bằng COHERE_RERANK_MODEL = rerank-v4.0-fast. Cách thiết kế này cho thấy reranker không thay thế retrieval, mà hoạt động như một lớp tinh chỉnh tùy chọn đặt sau retrieval và trước fusion.",
            ),
            (
                "Normal",
                "Ở tầng trả lời, OPENAI_LLM_MODEL hiện quyết định cả chế độ hybrid_rag, contextual_hybrid_rag và llm_only, còn OPENAI_EMBED_MODEL quyết định chất lượng vector retrieval. Vì vậy, khi thay đổi mô hình trong .env, người dùng không chỉ đổi output của câu trả lời mà còn gián tiếp đổi chất lượng candidate ở semantic branch. Đây là lý do nhóm trình bày phase 2 và phase 3 như hai lớp liên kết chặt chẽ về cấu hình, thay vì xem chúng là hai bước hoàn toàn độc lập.",
            ),
            (
                "Normal",
                "Ở tầng trả lời, chế độ hybrid_rag dùng evidence bundle từ hybrid retrieval để sinh grounded answer và trích dẫn nguồn. Chế độ contextual_hybrid_rag dùng evidence bundle từ contextual_hybrid retrieval, do đó kỳ vọng câu trả lời bám trọng tâm hơn khi context summary giúp evidence rõ nghĩa hơn. Chế độ llm_only bỏ qua retrieval và trả lời trực tiếp, đóng vai trò baseline closed-book.",
            ),
            (
                "Normal",
                "Điểm quan trọng trong chiến lược trả lời của repo là tách retrieval mode khỏi answer mode. Nhờ vậy, cùng một pipeline retrieval có thể được đánh giá độc lập trước khi ảnh hưởng của generation được đưa vào. Đây là lý do vì sao Chương 4 sẽ tách riêng kết quả retrieval, kết quả answer-level evaluation và pairwise comparison thay vì gộp tất cả thành một con số duy nhất.",
            ),
        ],
    }
)

SECTION_CONTENT.update(
    {
        "Mô hình không gian vector": [
            (
                "Normal",
                "Mô hình không gian vector biểu diễn mỗi tài liệu và mỗi truy vấn như một vector trong không gian nhiều chiều, trong đó mỗi chiều thường tương ứng với một term trong từ vựng. Nhờ cách biểu diễn này, truy hồi thông tin không còn là bài toán khớp điều kiện đúng-sai như Boolean retrieval, mà trở thành bài toán đo độ gần giữa truy vấn và tài liệu trong một không gian hình học.",
            ),
            (
                "Normal",
                "Trong trực giác của Vector Space Model, tài liệu càng chứa nhiều term quan trọng trùng với truy vấn thì vector của nó càng gần vector truy vấn. Điểm mạnh của mô hình là có thể xếp hạng theo mức độ liên quan thay vì chỉ lọc ra danh sách thỏa điều kiện. Đây cũng là nền tảng để hiểu vì sao các phương pháp lexical hiện đại như TF-IDF, cosine similarity và sau này là BM25 vẫn có quan hệ chặt chẽ với tư duy biểu diễn văn bản dưới dạng vector.",
            ),
            (
                "Normal",
                "Ở mức khái niệm, dense retrieval ngày nay cũng tiếp tục tinh thần của mô hình không gian vector, nhưng thay vì biểu diễn theo term bề mặt, hệ thống dùng embedding học được từ mô hình ngôn ngữ để ánh xạ truy vấn và tài liệu sang một không gian ngữ nghĩa liên tục. Vì vậy, việc nắm rõ VSM giúp liên kết tự nhiên giữa phần lý thuyết IR cổ điển và phần triển khai retrieval hiện đại trong repo.",
            ),
        ],
        "Dense retrieval, hybrid retrieval, RAG và LLM-only": [
            (
                "Normal",
                "Dense retrieval ánh xạ truy vấn và tài liệu vào cùng một không gian embedding, sau đó dùng một hàm khoảng cách như cosine similarity để tìm các vector gần nhau. So với lexical retrieval, dense retrieval mạnh hơn ở chỗ có thể bắt được quan hệ ngữ nghĩa ngay cả khi truy vấn và tài liệu không dùng cùng từ khóa bề mặt [2]. Tuy nhiên, nó cũng có thể bỏ sót các tín hiệu rất cụ thể như tên thuốc, viết tắt hoặc cụm từ chuyên ngành hiếm gặp.",
            ),
            (
                "Normal",
                "Vì sparse retrieval và dense retrieval có ưu điểm bổ sung cho nhau, nhiều hệ thống thực tế chọn hybrid retrieval. Trong repo hiện tại, hai nhánh BM25 và vector search được chạy song song rồi hợp nhất bằng Reciprocal Rank Fusion. Cách làm này giữ được độ nhạy với từ khóa rõ ràng của BM25, đồng thời tận dụng khả năng bắt ngữ nghĩa gần đúng của embedding retrieval. Ở phiên bản mới hơn, repo còn có contextual retrieval, trong đó mỗi chunk được gắn thêm context summary để cải thiện cả lexical matching lẫn semantic matching [12].",
            ),
            (
                "Normal",
                "Retrieval-Augmented Generation, hay RAG, là kiến trúc kết hợp retrieval với text generation [3]. Mô hình sinh không trả lời hoàn toàn dựa trên tri thức tham số, mà dựa trên evidence bundle vừa được truy hồi. Trong grounded mode, chất lượng câu trả lời phụ thuộc vào cả hai tầng: nếu retrieval bỏ sót tài liệu quan trọng thì generation khó có thể bù lại; ngược lại, nếu evidence đủ tốt, câu trả lời thường thận trọng và có khả năng truy vết tốt hơn.",
            ),
            (
                "Normal",
                "LLM-only là chế độ đối chứng trong đề tài. Ở chế độ này, mô hình chỉ nhìn thấy câu hỏi và trả lời trực tiếp theo kiểu closed-book. Đây là baseline cần thiết để kiểm tra xem việc thêm retrieval có thực sự cải thiện chất lượng câu trả lời hay không. Repo hiện vì thế không chỉ có cặp hybrid_rag và llm_only, mà còn có contextual_hybrid_rag như một grounded branch tăng cường. Ngoài ra, hệ thống còn hỗ trợ reranker Cohere ở mức tùy chọn để sắp lại candidate trước khi fuse, dù mặc định trong cấu hình mẫu tính năng này đang tắt.",
            ),
        ],
        "Các công trình nghiên cứu liên quan (ít nhất 5 công trình mới nhất)": [
            (
                "Normal",
                "Trong medical QA, một mốc quan trọng là Med-PaLM của Singhal và cộng sự [4]. Công trình này cho thấy LLM có thể đạt năng lực đáng kể trên nhiều bài toán y khoa, đồng thời nhấn mạnh rằng việc đánh giá mô hình cho miền y tế không thể chỉ dựa trên độ trôi chảy ngôn ngữ, mà phải xét cả độ đúng, độ an toàn và mức độ phù hợp với ngữ cảnh sử dụng.",
            ),
            (
                "Normal",
                "Nori và cộng sự [5] tiếp tục cho thấy GPT-4 có thể đạt kết quả cao trên các medical challenge problems. Tuy nhiên, kết quả tốt trên benchmark không đồng nghĩa với việc mô hình luôn trả lời tốt cho câu hỏi thực tế của người dùng. Điều này giải thích vì sao trong đồ án, llm_only chỉ được xem là baseline mạnh chứ không phải lời giải cuối cùng.",
            ),
            (
                "Normal",
                "Ở hướng xây dựng benchmark, Cai và cộng sự [6], Alonso và cộng sự [8], cùng Kim và cộng sự [9] đều góp phần mở rộng cách đánh giá medical QA theo hướng đa ngôn ngữ, đa dạng giải thích và nhiều kiểu câu hỏi hơn. Các benchmark này giúp nhìn rõ rằng chất lượng medical QA không chỉ nằm ở đáp án cuối, mà còn nằm ở khả năng đưa ra lý do và bằng chứng phù hợp.",
            ),
            (
                "Normal",
                "Ở hướng retrieval-augmented generation cho miền y khoa, Xiong và cộng sự [7] chỉ ra rằng việc thêm retrieval có thể cải thiện factual grounding, nhưng chất lượng cuối cùng vẫn phụ thuộc mạnh vào coverage của knowledge base và cách truy hồi evidence. Xiong và cộng sự [10] sau đó mở rộng thêm hướng iterative retrieval bằng câu hỏi follow-up, cho thấy medical RAG có thể được tăng cường thêm bằng nhiều bước truy hồi thay vì chỉ một lần truy hồi đơn giản.",
            ),
            (
                "Normal",
                "Liang và cộng sự [11] đề xuất RGAR theo hướng kết hợp retrieval với cơ chế sinh lặp, nhằm cải thiện tính factual-aware trong medical QA. Nhóm công trình này cho thấy xu hướng chung của lĩnh vực là không xem retrieval như một module phụ, mà xem nó là thành phần cốt lõi để kiểm soát độ tin cậy của câu trả lời.",
            ),
            (
                "Normal",
                "Bên cạnh các bài báo học thuật, một hướng engineering đáng chú ý là contextual retrieval của Anthropic [12]. Tài liệu này không phải benchmark y khoa, nhưng gợi ý một ý tưởng quan trọng: chất lượng retrieval có thể tăng nếu chunk không bị xem như một mảnh văn bản cô lập, mà được bổ sung ngữ cảnh tóm tắt từ tài liệu gốc. Repo hiện tại không sao chép nguyên xi thiết lập của Anthropic, nhưng đã hiện thực một biến thể thực dụng của ý tưởng này qua context_summary, contextualized_body và contextual_embedding.",
            ),
            (
                "Normal",
                "Từ các công trình trên có thể rút ra ba nhận xét. Thứ nhất, medical QA là bài toán mà độ đúng và độ thận trọng quan trọng không kém độ trôi chảy. Thứ hai, retrieval giúp grounding tốt hơn nhưng chỉ phát huy khi knowledge base và schema dữ liệu được thiết kế hợp lý. Thứ ba, các hướng contextual retrieval và iterative retrieval là xu hướng đáng chú ý cho các hệ thống IR/RAG thế hệ mới, và cũng là định hướng phát triển phù hợp cho repo hiện tại.",
            ),
        ],
        "Tổng quan phương pháp hiện thực": [
            (
                "Normal",
                "Hệ thống trong repo được tổ chức theo kiến trúc pipeline bốn pha, kèm theo một lớp demo app dùng để quan sát trực tiếp các artifact sinh ra trong quá trình chạy. Thiết kế này giúp tách biệt rõ trách nhiệm của từng bước, đồng thời giúp nhóm có thể chạy lại một pha riêng mà không phải làm lại toàn bộ hệ thống.",
            ),
            ("Normal", "Bốn giai đoạn cốt lõi lần lượt là:"),
            ("List Paragraph", "Chuẩn bị dữ liệu từ nhiều nguồn khác nhau và chuẩn hóa về schema chung."),
            ("List Paragraph", "Indexing để tách documents, nodes, contextual fields và chỉ mục PostgreSQL."),
            ("List Paragraph", "Retrieval và answer generation cho nhiều chế độ trả lời."),
            ("List Paragraph", "Evaluation để tính metric, tổng hợp system summary và failure cases."),
            (
                "Normal",
                "Điểm quan trọng của cách tổ chức này là repo không chỉ mô tả ý tưởng RAG ở mức khái niệm, mà hiện thực đủ cả vòng đời dữ liệu: từ crawl, chuẩn hóa, index, truy hồi, sinh câu trả lời, đánh giá đến trình bày qua giao diện web. Vì vậy, báo cáo cần bám theo luồng thực thi thật của repo thay vì chỉ mô tả một sơ đồ kiến trúc chung chung.",
            ),
        ],
    }
)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_before(target: Paragraph, text: str, style: str) -> Paragraph:
    new_para = target.insert_paragraph_before(text)
    new_para.style = style
    return new_para


def set_section_paragraphs(doc: Document, heading_text: str, items: SectionContent) -> None:
    paragraphs = doc.paragraphs
    start_idx = None
    end_idx = None
    possible_texts = {heading_text}
    if heading_text in HEADING_RENAMES:
        possible_texts.add(HEADING_RENAMES[heading_text][1])
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() in possible_texts and paragraph.style and paragraph.style.name.startswith("Heading"):
            start_idx = idx
            break
    if start_idx is None:
        raise ValueError(f"Heading not found: {heading_text}")

    for idx in range(start_idx + 1, len(paragraphs)):
        paragraph = paragraphs[idx]
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            end_idx = idx
            break
    if end_idx is None:
        end_idx = len(paragraphs)

    current = paragraphs[start_idx + 1 : end_idx]
    min_len = min(len(current), len(items))
    for idx in range(min_len):
        style_name, text = items[idx]
        current[idx].style = style_name
        current[idx].text = text

    for paragraph in current[min_len:]:
        delete_paragraph(paragraph)

    next_heading = None
    found_heading = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in possible_texts and paragraph.style and paragraph.style.name.startswith("Heading"):
            found_heading = True
            continue
        if found_heading and paragraph.style and paragraph.style.name.startswith("Heading"):
            next_heading = paragraph
            break
    if next_heading is None:
        raise ValueError(f"Next heading not found after: {heading_text}")

    extra = items[min_len:]
    for style_name, text in reversed(extra):
        insert_paragraph_before(next_heading, text, style_name)


def rename_headings(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        replacement = HEADING_RENAMES.get(text)
        if replacement is None:
            continue
        style_name, new_text = replacement
        paragraph.style = style_name
        paragraph.text = new_text


def ensure_column_count(table: Table, columns: int) -> None:
    while len(table.columns) < columns:
        table.add_column(Inches(1.6))


def ensure_row_count(table: Table, rows: int) -> None:
    while len(table.rows) < rows:
        table.add_row()


def fill_table(table: Table, rows: list[list[str]]) -> None:
    ensure_row_count(table, len(rows))
    ensure_column_count(table, len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value


def set_table_font_size(table: Table, size_pt: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)


def update_tables(doc: Document) -> None:
    fill_table(
        doc.tables[0],
        [
            ["Thành phần", "Vai trò trong hệ thống"],
            ["Python 3.13", "Ngôn ngữ chính để hiện thực pipeline dữ liệu, retrieval, evaluation và các script hỗ trợ báo cáo."],
            ["LlamaIndex", "Hỗ trợ embedding, wrapper mô hình và tích hợp retrieval/generation trong các pha xử lý."],
            ["OpenAI API", "Cung cấp embedding model, mô hình trả lời và evaluator dùng trong grounded answer."],
            ["PostgreSQL", "Lưu knowledge base, runtime logs và toàn bộ bảng phục vụ thực nghiệm."],
            ["pgvector", "Lưu vector embedding và tăng tốc semantic search bằng chỉ mục HNSW [13]."],
            ["pg_textsearch", "Hỗ trợ BM25 lexical retrieval trực tiếp trong PostgreSQL [14]."],
            ["FastAPI", "Cung cấp backend API cho demo app và bọc runtime Python thành HTTP endpoints."],
            ["Next.js 16", "Hiện thực frontend demo theo kiểu dashboard/workbench cục bộ."],
            ["React 19", "Xây dựng component hiển thị phase overview, evidence bundle, citation và failure cases."],
            ["Docker Compose", "Khởi tạo PostgreSQL cục bộ để bảo đảm tính tái hiện của môi trường chạy."],
        ],
    )
    fill_table(
        doc.tables[1],
        [
            ["Nguồn dữ liệu", "Số lượng", "Vai trò"],
            ["beir_nfcorpus", "2138 documents", "Backbone benchmark retrieval và nguồn qrels chính."],
            ["medlineplus_nutrition", "23 documents", "Nguồn public-health chính thống, tăng độ tin cậy của evidence."],
            ["fda_daily_value", "13 documents", "Bổ sung thông tin nhãn dinh dưỡng và giá trị khuyến nghị."],
            ["fda_label_pdf", "1 document", "Tài liệu tham chiếu PDF về Nutrition Facts."],
            ["pubmed_nutrition", "98 documents", "Bổ sung góc nhìn nghiên cứu và diversity cho evidence."],
            ["Tổng document", "2273", "Tập tài liệu đầu vào sau chuẩn hóa và hợp nhất."],
            ["Tổng node", "2320", "Đơn vị truy hồi thực tế sau bước chunking."],
        ],
    )
    fill_table(
        doc.tables[2],
        [
            ["Nhóm schema", "Trường tiêu biểu", "Vai trò"],
            ["Document-level", "doc_id, source_id, title, source_url, document_text, document_token_count, chunk_count, section_type", "Mô tả tài liệu gốc và mức độ chia chunk trong knowledge base."],
            ["Node-level", "node_id, doc_id, body, raw_body, context_summary, contextualized_body, chunk_index, chunk_count, token_count, char_count, embedding, contextual_embedding", "Đơn vị truy hồi phục vụ lexical search, vector search và contextual retrieval."],
            ["Experiment-level", "retrieval_runs, answer_runs, answer_evaluations, comparison_runs", "Lưu dấu vết chạy, kết quả trả lời, điểm chấm và so sánh pairwise."],
        ],
    )
    fill_table(
        doc.tables[3],
        [
            ["Chỉ số", "hybrid", "contextual_hybrid"],
            ["recall@10", "0.195429", "0.193321"],
            ["MRR@10", "0.548713", "0.531045"],
            ["nDCG@10", "0.353326", "0.341681"],
            ["MAP", "0.158659", "0.155346"],
            ["unique / duplicate@10", "-", "9.9023 / 0.0977"],
        ],
    )
    fill_table(
        doc.tables[4],
        [
            ["Chế độ", "Faith.", "Correct.", "Relev.", "Mẫu/Ghi chú"],
            ["Hybrid RAG", "0.9548", "0.8224", "0.8876", "307 mẫu"],
            ["Contextual Hybrid RAG", "0.9493", "0.8284", "0.9216", "307 mẫu"],
            ["LLM-only", "0.0000", "0.4243", "0.9911", "307 mẫu"],
        ],
    )
    fill_table(
        doc.tables[5],
        [
            ["So sánh", "Left wins", "Left win rate", "Trạng thái"],
            ["hybrid_rag vs llm_only", "223/307", "0.7264", "clean"],
            ["contextual_hybrid_rag vs hybrid_rag", "276/307", "0.8990", "clean"],
            ["contextual_hybrid_rag vs llm_only", "169/307", "0.5505", "72 rows quota-failed"],
            ["Ghi chú", "-", "-", "Cặp contextual vs llm_only chỉ nên xem là tham khảo"],
        ],
    )
    set_table_font_size(doc.tables[0], 10)
    set_table_font_size(doc.tables[1], 10)
    set_table_font_size(doc.tables[2], 10)
    set_table_font_size(doc.tables[3], 11)
    set_table_font_size(doc.tables[4], 10)
    set_table_font_size(doc.tables[5], 10)


def update_references(doc: Document) -> None:
    heading = "TÀI LIỆU THAM KHẢO"
    paragraphs = doc.paragraphs
    start_idx = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == heading and paragraph.style and paragraph.style.name.startswith("Heading"):
            start_idx = idx
            break
    if start_idx is None:
        raise ValueError("Reference heading not found")

    current = doc.paragraphs[start_idx + 1 :]
    min_len = min(len(current), len(REFERENCE_LINES))
    for idx in range(min_len):
        current[idx].style = "Normal"
        current[idx].text = REFERENCE_LINES[idx]
    for paragraph in current[min_len:]:
        delete_paragraph(paragraph)
    for line in REFERENCE_LINES[min_len:]:
        doc.add_paragraph(line, style="Normal")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    source_path = DOCX_PATH
    if not DOCX_PATH.exists():
        source_path = UPDATED_COPY_PATH

    doc = Document(source_path)
    for heading, items in SECTION_CONTENT.items():
        set_section_paragraphs(doc, heading, items)
    rename_headings(doc)
    update_tables(doc)
    update_references(doc)
    save_errors: list[str] = []
    saved_main = False
    try:
        doc.save(DOCX_PATH)
        saved_main = True
    except PermissionError as exc:
        save_errors.append(str(exc))
    try:
        doc.save(UPDATED_COPY_PATH)
        final_copy = UPDATED_COPY_PATH
    except PermissionError:
        doc.save(FALLBACK_COPY_PATH)
        final_copy = FALLBACK_COPY_PATH
    if saved_main:
        print(f"Updated: {DOCX_PATH}")
    else:
        print(f"Main file locked; updated copy instead: {final_copy}")
        for err in save_errors:
            print(f"Lock info: {err}")
    print(f"Backup: {BACKUP_PATH}")
    print(f"Copy: {final_copy}")


if __name__ == "__main__":
    main()
